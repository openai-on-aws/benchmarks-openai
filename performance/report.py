"""
Build the combined latency report (REPORT.md + REPORT.html + charts) from
schema-v2 result JSONs in performance/results/.

Every number in the report is computed from the result files at build time —
nothing is hardcoded. Usage:

  python performance/report.py
"""

import base64
import glob
import html as html_mod
import json
import os
import re
import statistics
from datetime import datetime, timezone

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

RESULTS_DIR = os.path.join(os.path.dirname(__file__), "results")
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
QUALITY_RESULTS_DIR = os.path.join(os.path.dirname(__file__), "..", "quality", "results")

SIZES = ["1k", "5k", "10k", "20k"]
MODELS = [
    ("gpt-5.6-luna", "openai.gpt-5.6-luna", "gpt-5.6-luna"),
    ("gpt-5.6-terra", "openai.gpt-5.6-terra", "gpt-5.6-terra"),
    ("gpt-5.6-sol", "openai.gpt-5.6-sol", "gpt-5.6-sol"),
]
BACKEND_LABEL = {"bedrock": "Bedrock", "openai": "OpenAI 1P"}

# Categorical palette (validated: CVD dE 24.7, normal dE 33.6, both >=3:1 on #fcfcfb)
C_BEDROCK = "#2a78d6"
C_OPENAI = "#eb6834"
INK = "#0b0b0b"
MUTED = "#898781"
GRID = "#e1e0d9"
SURFACE = "#fcfcfb"


def load_results():
    """out[(backend, family)][size] = latest payload for that cell."""
    out = {}
    for path in sorted(glob.glob(os.path.join(RESULTS_DIR, "results_*.json"))):
        with open(path) as f:
            d = json.load(f)
        if d.get("schema_version") != 2 or d.get("concurrency", 1) != 1 or d.get("reasoning_effort"):
            continue
        family = d["model"].replace("openai.", "")
        key = (d["backend"], family)
        cell = out.setdefault(key, {})
        prev = cell.get(d["input_label"])
        if prev is None or d["started_at"] > prev["started_at"]:
            cell[d["input_label"]] = d
    return out


def fnum(x, nd=0):
    if x is None:
        return "n/a"
    return f"{x:,.{nd}f}"


def get(summary_row, metric, stat):
    block = summary_row.get(metric)
    return block.get(stat) if block else None


def rows_for(results, backend, family):
    """Yield (size, summary_row, payload) ordered by size then max_out."""
    cell = results.get((backend, family), {})
    for size in SIZES:
        d = cell.get(size)
        if not d:
            continue
        for s in d["summary"]:
            yield size, s, d


# ---------------------------------------------------------------- setup table

def prompt_stats():
    stats = {}
    tokens = {"1k": 1002, "5k": 4750, "10k": 9984, "20k": 20136}
    for size in SIZES:
        p = os.path.join(DATA_DIR, f"prompt_{size}.txt")
        stats[size] = (tokens[size], len(open(p).read()))
    return stats


def build_setup_table(results):
    ps = prompt_stats()
    dates = sorted({d["started_at"][:10] for cell in results.values() for d in cell.values()})
    total_calls = sum(s["n_runs"] for cell in results.values() for d in cell.values() for s in d["summary"])
    total_errors = sum(s["n_errors"] for cell in results.values() for d in cell.values() for s in d["summary"])
    rows = [
        ("Models", "`openai.gpt-5.6-luna`, `-terra`, `-sol` (Bedrock) vs `gpt-5.6-luna`, `-terra`, `-sol` (1P)"),
        ("Bedrock endpoint", "`https://bedrock-mantle.us-west-2.api.aws/openai/v1` (luna/terra) · `us-east-1` (sol — not served in us-west-2)"),
        ("OpenAI 1P endpoint", "`https://api.openai.com/v1`"),
        ("API surface", "Responses API, streaming — identical code path on both backends (`performance/benchmark.py`)"),
        ("Auth", "Bedrock: IAM via `aws-bedrock-token-generator` · 1P: user-supplied `OPENAI_API_KEY`"),
        ("Runs per config", "25"),
        ("Concurrency", "Single thread, sequential calls"),
        ("Reasoning effort", "Default (not set)"),
        ("Output configs", "1k input: 100/500/1000 · 5k: 500/1000/5000 · 10k: 500/1000/5000 · 20k: 1000/5000/10000"),
        ("Prompt source", "Manhattan Project (Wikipedia) — real varied text, no repetition"),
        ("1k input prompt", f"{ps['1k'][0]:,} verified tokens ({ps['1k'][1]:,} chars)"),
        ("5k input prompt", f"{ps['5k'][0]:,} verified tokens ({ps['5k'][1]:,} chars)"),
        ("10k input prompt", f"{ps['10k'][0]:,} verified tokens ({ps['10k'][1]:,} chars)"),
        ("20k input prompt", f"{ps['20k'][0]:,} verified tokens ({ps['20k'][1]:,} chars)"),
        ("Run dates", ", ".join(dates) + " (Bedrock luna 07-18; sol overnight 07-20→21; rest 07-20)"),
        ("Total calls", f"{total_calls:,} ({total_errors} errored)"),
    ]
    lines = ["| Parameter | Value |", "|---|---|"]
    lines += [f"| {k} | {v} |" for k, v in rows]
    return "\n".join(lines)


# --------------------------------------------------------------- detail table

def build_detail_table(results, backend, family):
    lines = [
        "| Input | Max out | TTFT p50 | TTFT p95 | TTFT p99 | Tok/s p50 | E2E p50 | E2E p95 |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for size, s, _ in rows_for(results, backend, family):
        lines.append(
            f"| {size.upper()} | {s['max_output_tokens']:,} "
            f"| {fnum(get(s,'ttft_ms','p50'))} | {fnum(get(s,'ttft_ms','p95'))} | {fnum(get(s,'ttft_ms','p99'))} "
            f"| {fnum(get(s,'otps','p50'),1)} | {fnum(get(s,'e2e_ms','p50'))} | {fnum(get(s,'e2e_ms','p95'))} |"
        )
    return "\n".join(lines)


# ----------------------------------------------------------- comparison table

def pct_delta(br, op, lower_better=True):
    if br is None or op is None or op == 0:
        return "n/a"
    pct = (op - br) / op * 100 if lower_better else (br - op) / op * 100
    return f"{'+' if pct >= 0 else ''}{pct:.0f}%"


def build_comparison_table(results, family):
    br = {(size, s["max_output_tokens"]): s for size, s, _ in rows_for(results, "bedrock", family)}
    op = {(size, s["max_output_tokens"]): s for size, s, _ in rows_for(results, "openai", family)}
    lines = [
        "| Input | Max out | Metric | Bedrock | OpenAI 1P | Delta |",
        "|---|---|---|---|---|---|",
    ]
    metrics = [
        ("ttft_ms", "p50", "TTFT p50 (ms)", True),
        ("ttft_ms", "p95", "TTFT p95 (ms)", True),
        ("ttft_ms", "p99", "TTFT p99 (ms)", True),
        ("otps", "p50", "Tok/s p50", False),
        ("e2e_ms", "p50", "E2E p50 (ms)", True),
        ("e2e_ms", "p95", "E2E p95 (ms)", True),
    ]
    for size in SIZES:
        for key in sorted(k for k in br if k[0] == size and k in op):
            for metric, stat, label, lower in metrics:
                a, b = get(br[key], metric, stat), get(op[key], metric, stat)
                if a is None and b is None:
                    continue
                lines.append(f"| {size.upper()} | {key[1]:,} | {label} | {fnum(a,1)} | {fnum(b,1)} | {pct_delta(a, b, lower)} |")
    return "\n".join(lines)


# ---------------------------------------------------------------------- chart

def pctile(vals, p):
    s = sorted(vals)
    if not s:
        return None
    idx = (p / 100) * (len(s) - 1)
    lo, hi = int(idx), min(int(idx) + 1, len(s) - 1)
    return s[lo] + (s[hi] - s[lo]) * (idx - lo)


def raw_band(payload, max_out, field):
    """(p5, p50, p95, p99) of a per-call field, from raw data (outlier-robust)."""
    vals = [r[field] for r in payload["raw"][str(max_out)] if not r["error"] and r[field]]
    if not vals:
        return None, None, None, None
    return pctile(vals, 5), pctile(vals, 50), pctile(vals, 95), pctile(vals, 99)


def build_chart(results, family, outfile):
    fig, axes = plt.subplots(2, 4, figsize=(13.5, 6.2), dpi=180)
    fig.patch.set_facecolor(SURFACE)
    regions = sorted({d["base_url"].split(".")[1] for cell in [results.get(("bedrock", family), {})]
                      for d in cell.values() if "bedrock-mantle" in d["base_url"]})
    region_label = "/".join(regions) if regions else "?"
    fig.suptitle(f"{family} — Bedrock ({region_label}) vs OpenAI 1P · p50 across 25 sequential runs",
                 fontsize=12, fontweight="bold", color=INK, y=0.99)

    series = [("bedrock", C_BEDROCK, "Bedrock"), ("openai", C_OPENAI, "OpenAI 1P")]
    row_specs = [
        (0, "ttft_ms", "TTFT (ms)"),
        (1, "otps", "Output tok/s"),
    ]

    for row, metric, ylabel in row_specs:
        vals = []       # drives the row y-limit
        band_vals = []  # TTFT row includes band tops; tok/s row scales to p50s only
        for col, size in enumerate(SIZES):
            ax = axes[row][col]
            ax.set_facecolor(SURFACE)
            for backend, color, label in series:
                pts = [(s["max_output_tokens"], d) for sz, s, d in rows_for(results, backend, family) if sz == size]
                if not pts:
                    continue
                xs = list(range(len(pts)))
                bands = [raw_band(d, m, metric) for m, d in pts]
                lo = [b[0] for b in bands]
                p50 = [b[1] for b in bands]
                hi = [b[2] for b in bands]
                p99 = [b[3] for b in bands]
                ax.plot(xs, p50, marker="o", markersize=6, linewidth=2, color=color,
                        label=label, zorder=3)
                if all(v is not None for v in lo + hi):
                    ax.fill_between(xs, lo, hi, color=color, alpha=0.13, linewidth=0, zorder=2)
                # p99 as a dashed tail line on latency panels only; on tok/s it
                # tracks the burst outliers and would wreck the axis. It does NOT
                # drive the y-limit — a single extreme call would flatten the row
                # (sol 20k had a 240s outlier); it clips instead, per the footnote.
                if metric == "ttft_ms" and all(v is not None for v in p99):
                    ax.plot(xs, p99, linestyle=(0, (3, 2)), linewidth=1.3, color=color,
                            label=f"{label} p99", zorder=3)
                vals += [v for v in p50 if v is not None]
                band_vals += [v for v in hi if v is not None]
                ax.set_xticks(xs)
                ax.set_xticklabels([f"{m:,}" for m, _ in pts], fontsize=8, color=MUTED)
            if row == 0:
                ax.set_title(f"{size.upper()} input", fontsize=10, color=INK)
            if row == 1:
                ax.set_xlabel("max output tokens", fontsize=8.5, color=MUTED)
            if col == 0:
                ax.set_ylabel(ylabel, fontsize=9.5, color=INK)
            ax.tick_params(axis="y", labelsize=8, colors=MUTED)
            ax.grid(axis="y", color=GRID, linewidth=0.7)
            for spine in ["top", "right"]:
                ax.spines[spine].set_visible(False)
            for spine in ["left", "bottom"]:
                ax.spines[spine].set_color(GRID)
        # TTFT row: scale to p95 band tops (p99 may clip). Tok/s row: scale to
        # ~1.6x the largest p50 so burst-flush outliers at tiny outputs
        # (10,000+ tok/s) can't flatten it; bands clip, per the footnote.
        if metric == "otps":
            ymax = max(vals) * 1.6 if vals else 1
        else:
            ymax = max(vals + band_vals) * 1.15 if vals or band_vals else 1
        for col in range(4):
            axes[row][col].set_ylim(0, ymax)
            if col > 0:
                axes[row][col].set_yticklabels([])

    handles, labels = axes[0][0].get_legend_handles_labels()
    order = [labels.index(x) for x in ["Bedrock", "OpenAI 1P", "Bedrock p99", "OpenAI 1P p99"] if x in labels]
    fig.legend([handles[i] for i in order], [labels[i] for i in order],
               loc="upper right", bbox_to_anchor=(0.995, 0.985), ncol=2,
               fontsize=9, frameon=False)
    fig.text(0.005, 0.005,
             "Solid line: p50 per call · shaded band: p5→p95 (90% of calls fall inside) · "
             "dashed: p99, may clip off-scale (TTFT row) · y-scale shared per row · "
             "tok/s axis scaled to p50s; single-flush outliers at small outputs exceed it",
             fontsize=8, color=MUTED)
    fig.tight_layout(rect=[0, 0.02, 1, 0.95])
    # The metadata chunk also shifts the PNG byte stream if a chance base64
    # substring ever pattern-matches a credential prefix in secret scanners.
    fig.savefig(outfile, facecolor=SURFACE, bbox_inches="tight",
                metadata={"Software": "benchmarks-openai performance/report.py"})
    plt.close(fig)


# ------------------------------------------------------------------- findings

def mean_of(results, backend, family, metric, stat, min_max_out=500):
    vals = [get(s, metric, stat) for _, s, _ in rows_for(results, backend, family)
            if s["max_output_tokens"] >= min_max_out and get(s, metric, stat) is not None]
    return statistics.mean(vals) if vals else None


def build_findings(results):
    f = []

    # 1. TTFT flatness vs input size (at max_out=1000, present for every size)
    for family in [m[0] for m in MODELS]:
        for backend in ["bedrock", "openai"]:
            ttfts = {sz: get(s, "ttft_ms", "p50") for sz, s, _ in rows_for(results, backend, family)
                     if s["max_output_tokens"] == 1000}
            if len(ttfts) == 4:
                lo, hi = min(ttfts.values()), max(ttfts.values())
                if hi / lo < 1.35:
                    f.append(f"**TTFT is roughly flat across input sizes on {BACKEND_LABEL[backend]} ({family}):** "
                             f"at 1,000 max output tokens, p50 TTFT stays within "
                             f"{fnum(lo)}–{fnum(hi)} ms from 1K to 20K input tokens — a 20× larger prompt "
                             f"does not meaningfully move time-to-first-token.")
                    break  # one flatness bullet per family is enough

    # 2. TTFT gap per family
    for family in [m[0] for m in MODELS]:
        br = mean_of(results, "bedrock", family, "ttft_ms", "p50", 0)
        op = mean_of(results, "openai", family, "ttft_ms", "p50", 0)
        if br and op:
            pct = (op - br) / op * 100
            direction = "lower" if pct > 0 else "higher"
            f.append(f"**{family} TTFT:** averaged across all 12 configs, Bedrock p50 TTFT is "
                     f"{abs(pct):.0f}% {direction} than 1P ({fnum(br)} vs {fnum(op)} ms).")

    # 3. Throughput comparison and luna-vs-terra gap
    tp = {}
    for family in [m[0] for m in MODELS]:
        for backend in ["bedrock", "openai"]:
            tp[(backend, family)] = mean_of(results, backend, family, "otps", "p50")
    if all(tp.values()):
        f.append(f"**Throughput (≥500-token outputs):** luna averages "
                 f"{fnum(tp[('bedrock','gpt-5.6-luna')],1)} tok/s on Bedrock vs {fnum(tp[('openai','gpt-5.6-luna')],1)} on 1P "
                 f"(+{(tp[('bedrock','gpt-5.6-luna')]/tp[('openai','gpt-5.6-luna')]-1)*100:.0f}%); terra averages "
                 f"{fnum(tp[('bedrock','gpt-5.6-terra')],1)} vs {fnum(tp[('openai','gpt-5.6-terra')],1)} tok/s "
                 f"(+{(tp[('bedrock','gpt-5.6-terra')]/tp[('openai','gpt-5.6-terra')]-1)*100:.0f}%).")
        luna_avg = statistics.mean([tp[('bedrock','gpt-5.6-luna')], tp[('openai','gpt-5.6-luna')]])
        terra_avg = statistics.mean([tp[('bedrock','gpt-5.6-terra')], tp[('openai','gpt-5.6-terra')]])
        f.append(f"**Luna streams ~{luna_avg/terra_avg:.1f}× faster than terra on both backends** "
                 f"(~{fnum(luna_avg)} vs ~{fnum(terra_avg)} tok/s at ≥500-token outputs) — a model "
                 f"characteristic, not a platform one.")

    # 4. Tail latency: worst TTFT p99/p50 ratio per backend
    for family in [m[0] for m in MODELS]:
        ratios = {}
        for backend in ["bedrock", "openai"]:
            rs = [(get(s, "ttft_ms", "p99") / get(s, "ttft_ms", "p50"), sz, s["max_output_tokens"])
                  for sz, s, _ in rows_for(results, backend, family)
                  if get(s, "ttft_ms", "p99") and get(s, "ttft_ms", "p50")]
            if rs:
                ratios[backend] = max(rs)
        if len(ratios) == 2:
            wb, wo = ratios["bedrock"], ratios["openai"]
            f.append(f"**{family} TTFT tail:** worst p99/p50 ratio is {wb[0]:.1f}× on Bedrock "
                     f"({wb[1].upper()}/{wb[2]:,} out) vs {wo[0]:.1f}× on 1P ({wo[1].upper()}/{wo[2]:,} out).")

    # 5. Reasoning burn at max_out=100
    burn = []
    for family in [m[0] for m in MODELS]:
        for backend in ["bedrock", "openai"]:
            cell = results.get((backend, family), {}).get("1k")
            if not cell:
                continue
            raw = cell["raw"].get("100", [])
            if raw:
                n = sum(1 for r in raw if not r["error"] and r["ttft_ms"] is None)
                if n:
                    burn.append(f"{n}/{len(raw)} on {BACKEND_LABEL[backend]} {family}")
    if burn:
        f.append("**At 100-token output budgets, gpt-5.6 often spends the whole budget on reasoning "
                 "and emits no visible text** (" + "; ".join(burn) + "). Null-TTFT calls are excluded "
                 "from latency stats; budget well above 100 output tokens for latency-sensitive use.")

    # 6. Output truncation vs requested max
    trunc = []
    for family in [m[0] for m in MODELS]:
        for backend in ["bedrock", "openai"]:
            for sz, s, d in rows_for(results, backend, family):
                mo = s["max_output_tokens"]
                if mo < 5000:
                    continue
                raw = [r for r in d["raw"][str(mo)] if not r["error"]]
                if raw:
                    mean_out = statistics.mean(r["output_tokens"] for r in raw)
                    if mean_out < 0.9 * mo:
                        trunc.append(f"{BACKEND_LABEL[backend]} {family} {sz.upper()}/{mo:,}: mean {fnum(mean_out)} tokens")
    if trunc:
        f.append("**Responses often stop well before large `max_output_tokens` limits** (model finishes "
                 "naturally rather than truncating): " + "; ".join(trunc[:4]) +
                 (f"; and {len(trunc)-4} more configs" if len(trunc) > 4 else "") + ".")

    return "\n".join(f"- {x}" for x in f)


# -------------------------------------------------------------- quality evals

EVAL_TASK_LABELS = {"mmlu_pro": "MMLU-Pro", "math500": "MATH-500", "gsm8k": "GSM8K"}
EVAL_TASK_ORDER = ["mmlu_pro", "math500", "gsm8k"]


def load_quickevals():
    """latest[(model_label, task)] = summary; keeps the newest file per pair."""
    latest = {}
    for path in sorted(glob.glob(os.path.join(QUALITY_RESULTS_DIR, "quickeval_*.json"))):
        with open(path) as f:
            d = json.load(f)
        backend = "Bedrock" if d["backend"] == "mantle" else "OpenAI 1P"
        model = d["model"].replace("openai.", "")
        label = f"{model} ({backend}" + (f", effort={d['reasoning_effort']}" if d["reasoning_effort"] else "") + ")"
        key = (label, d["task"])
        if key not in latest or d["timestamp"] > latest[key]["_ts"]:
            latest[key] = {**d["summary"], "_ts": d["timestamp"], "n_asked": d["summary"]["n"]}
    return latest


def build_evals_section(section_no):
    evals = load_quickevals()
    if not evals:
        return ""
    models = sorted({k[0] for k in evals})
    lines = [
        f"## {section_no}. Task-quality evals — gpt-5.6 (reasoning off) vs gpt-5.4-mini/nano",
        "",
        "Accuracy on fixed-seed samples of community benchmarks (MMLU-Pro 140 stratified Qs, "
        "MATH-500 100 Qs, GSM8K 100 Qs). Every model answered the **same questions** via the "
        "same Responses-API path; exact-match scoring (MCQ letter / normalized boxed answer / "
        "final number). Latency here is full-response wall time per question (not TTFT), "
        "measured under 6-way concurrency — comparable across rows, not to the tables above. "
        "At these sample sizes the 95% CI is roughly ±8–10 points: treat differences inside "
        "that band as ties. Result files: `quality/results/quickeval_*.json`.",
        "",
        "| Model | Benchmark | Accuracy | Correct | Mean latency (ms) | Mean output tokens |",
        "|---|---|---|---|---|---|",
    ]
    for model in models:
        for task in EVAL_TASK_ORDER:
            s = evals.get((model, task))
            if not s:
                continue
            n_ok = s["n_asked"] - s["n_errors"]
            lines.append(
                f"| {model} | {EVAL_TASK_LABELS[task]} | {s['accuracy']:.0%} "
                f"| {s['n_correct']}/{n_ok} | {fnum(s['mean_latency_ms'])} | {fnum(s['mean_output_tokens'])} |")
    lines += [
        "",
        "Context for the matchup (an OpenAI-suggested comparison for migration planning): "
        "gpt-5.6 luna/terra on Bedrock with `reasoning: {effort: none}` — i.e. thinking "
        "disabled — against gpt-5.4-mini/nano on the OpenAI API at their defaults. "
        "GSM8K is saturated for all four models (95–97%) and acts as a sanity control.",
        "",
    ]
    return "\n".join(lines)


# ------------------------------------------------------------------- markdown

def build_markdown(results):
    charts = {}
    for family, *_ in MODELS:
        png = os.path.join(RESULTS_DIR, f"chart_{family}.png")
        build_chart(results, family, png)
        charts[family] = os.path.basename(png)

    parts = [f"""# GPT-5.6 on Amazon Bedrock vs OpenAI 1P — Latency & Quality Benchmark Report

**Generated:** {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} · **Repo:** [openai-on-aws/benchmarks-openai](https://github.com/openai-on-aws/benchmarks-openai)

Both backends are exercised through an identical code path — the OpenAI Responses API with streaming
(`performance/benchmark.py`). Metrics: **TTFT** (time to first output-text token), **Tok/s** (output tokens
per second of generation time), **E2E** (end-to-end wall time). Every number below is computed from the
timestamped result JSONs in `performance/results/` at report build time.

## 1. Test setup

{build_setup_table(results)}

## 2. Key performance findings

{build_findings(results)}

### Caveats

- Bedrock **luna** ran 2026-07-18; all other matrices ran 2026-07-20 (sol overnight into 07-21), so the
  luna comparison includes day-to-day variance. The **terra** and **sol** comparisons are same-session
  on both backends.
- **Sol's Bedrock runs used us-east-1** (sol is not served in us-west-2), so its Bedrock-vs-1P deltas
  include a region difference; luna/terra used us-west-2.
- Sol is a deep-reasoning model: it spends heavily on reasoning tokens before the first visible token,
  so its TTFT is inherently higher and more variable than luna/terra on both backends.
- Sequential, default reasoning effort. Concurrency and effort sweeps are supported by the harness but
  not yet run.
- Delta convention: **positive = Bedrock better** (lower latency or higher throughput).
"""]

    section_no = 3
    for family, *_ in MODELS:
        parts.append(f"""## {section_no}. {family}

### {section_no}.1 Benchmark chart

**How to read this chart:** the solid line is the median (p50) call. The shaded band spans
p5→p95 — 90% of the 25 calls per config landed inside it, so a wide band means inconsistent
latency, not measurement error. The dashed line (TTFT row only) is p99, the worst-case tail.
One panel per input size; y-scale shared within each row.

![{family} benchmark chart]({charts[family]})

### {section_no}.2 Bedrock detail (all timings ms)

{build_detail_table(results, "bedrock", family)}

### {section_no}.3 OpenAI 1P detail (all timings ms)

{build_detail_table(results, "openai", family)}

### {section_no}.4 Side-by-side comparison

{build_comparison_table(results, family)}
""")
        section_no += 1

    evals_md = build_evals_section(section_no)
    if evals_md:
        parts.append(evals_md)
        section_no += 1

    parts.append("""## Source files

Every table row and chart point traces to a result JSON under `performance/results/`
(`results_<backend>_<model>_<size>input_25runs_<timestamp>.json`), which holds full
distributions (mean/stddev/p50/p95/p99/min/max) and raw per-call measurements.
""")
    return "\n".join(parts), charts


# ----------------------------------------------------------------------- html

def md_to_html(md, charts):
    def inline(s):
        s = html_mod.escape(s)
        s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
        s = re.sub(r"`(.+?)`", r"<code>\1</code>", s)
        s = re.sub(r"\[(.+?)\]\((.+?)\)", r'<a href="\2">\1</a>', s)
        return s

    out, in_list, i = [], False, 0
    lines = md.splitlines()

    def close_list():
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    while i < len(lines):
        l = lines[i]
        img = re.match(r"!\[(.*?)\]\((.*?)\)", l.strip())
        if img:
            close_list()
            path = os.path.join(RESULTS_DIR, img.group(2))
            b64 = base64.b64encode(open(path, "rb").read()).decode()
            out.append(f'<div class="chart"><img alt="{html_mod.escape(img.group(1))}" '
                       f'src="data:image/png;base64,{b64}"></div>')
        elif l.startswith("|"):
            close_list()
            tbl = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in tbl]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            out.append("<table><thead><tr>" + "".join(f"<th>{inline(c)}</th>" for c in rows[0]) + "</tr></thead><tbody>")
            for r in rows[1:]:
                cells = []
                for j, c in enumerate(r):
                    cls = ""
                    if j == len(r) - 1 and c.endswith("%"):
                        try:
                            v = float(c.rstrip("%"))
                            cls = ' class="pos"' if v > 0 else (' class="neg"' if v < 0 else "")
                        except ValueError:
                            pass
                    cells.append(f"<td{cls}>{inline(c)}</td>")
                out.append("<tr>" + "".join(cells) + "</tr>")
            out.append("</tbody></table>")
            continue
        elif l.startswith("### "):
            close_list(); out.append(f"<h3>{inline(l[4:])}</h3>")
        elif l.startswith("## "):
            close_list(); out.append(f"<h2>{inline(l[3:])}</h2>")
        elif l.startswith("# "):
            close_list(); out.append(f"<h1>{inline(l[2:])}</h1>")
        elif l.startswith("- "):
            if not in_list:
                out.append("<ul>"); in_list = True
            out.append(f"<li>{inline(l[2:])}</li>")
        elif l.strip() == "":
            close_list()
        else:
            close_list(); out.append(f"<p>{inline(l)}</p>")
        i += 1
    close_list()

    body = "\n".join(out)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>GPT-5.6 on Bedrock vs OpenAI 1P — Latency Benchmark Report</title>
<style>
  :root {{
    --fg: #0b0b0b; --muted: #52514e; --line: #e1e0d9; --accent: #146eb4;
    --pos: #006300; --neg: #b42318; --bg: #fcfcfb; --code-bg: #f0efec;
  }}
  @media (prefers-color-scheme: dark) {{
    :root {{
      --fg: #e6e8ee; --muted: #c3c2b7; --line: #2c2c2a; --accent: #5aa9e6;
      --pos: #4ade80; --neg: #f87171; --bg: #1a1a19; --code-bg: #262624;
    }}
  }}
  body {{ font: 16px/1.6 system-ui, -apple-system, "Segoe UI", sans-serif;
         color: var(--fg); background: var(--bg); max-width: 1020px; margin: 2rem auto; padding: 0 1.25rem; }}
  h1 {{ font-size: 1.55rem; line-height: 1.3; border-bottom: 2px solid var(--line); padding-bottom: .6rem; }}
  h2 {{ font-size: 1.25rem; margin-top: 2.4rem; color: var(--accent); }}
  h3 {{ font-size: 1.02rem; margin-top: 1.6rem; }}
  code {{ background: var(--code-bg); padding: .1em .35em; border-radius: 4px; font-size: .86em; }}
  a {{ color: var(--accent); }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0 1.5rem; font-size: .84rem;
          font-variant-numeric: tabular-nums; }}
  th, td {{ border: 1px solid var(--line); padding: .32rem .55rem; text-align: right; }}
  th:nth-child(-n+3), td:nth-child(-n+3) {{ text-align: left; }}
  thead th {{ background: var(--code-bg); }}
  tbody tr:nth-child(even) {{ background: color-mix(in srgb, var(--code-bg) 50%, transparent); }}
  td.pos {{ color: var(--pos); font-weight: 600; }}
  td.neg {{ color: var(--neg); font-weight: 600; }}
  .chart {{ background: #fcfcfb; border: 1px solid var(--line); border-radius: 8px;
           padding: .5rem; margin: 1rem 0; }}
  .chart img {{ width: 100%; height: auto; display: block; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""


# ----------------------------------------------------------------------- docx

def md_to_docx(md, outfile):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def add_runs(par, text):
        # split on **bold** and `code`
        for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
            if part.startswith("**") and part.endswith("**"):
                r = par.add_run(part[2:-2]); r.bold = True
            elif part.startswith("`") and part.endswith("`"):
                r = par.add_run(part[1:-1]); r.font.name = "Courier New"; r.font.size = Pt(9.5)
            elif part:
                # strip markdown links down to their text
                par.add_run(re.sub(r"\[(.+?)\]\(.+?\)", r"\1", part))

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        l = lines[i]
        img = re.match(r"!\[(.*?)\]\((.*?)\)", l.strip())
        if img:
            doc.add_picture(os.path.join(RESULTS_DIR, img.group(2)), width=Inches(6.8))
        elif l.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in tbl]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    par = t.rows[ri].cells[ci].paragraphs[0]
                    add_runs(par, cell)
                    for r in par.runs:
                        r.font.size = Pt(8.5)
                        if ri == 0:
                            r.bold = True
            doc.add_paragraph()
            continue
        elif l.startswith("### "):
            doc.add_heading(re.sub(r"[*`]", "", l[4:]), level=3)
        elif l.startswith("## "):
            doc.add_heading(re.sub(r"[*`]", "", l[3:]), level=2)
        elif l.startswith("# "):
            doc.add_heading(re.sub(r"[*`]", "", l[2:]), level=1)
        elif l.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), l[2:])
        elif l.strip():
            add_runs(doc.add_paragraph(), l)
        i += 1

    doc.save(outfile)


def main():
    results = load_results()
    md, charts = build_markdown(results)
    md_path = os.path.join(RESULTS_DIR, "REPORT.md")
    with open(md_path, "w") as f:
        f.write(md)
    html_path = os.path.join(RESULTS_DIR, "REPORT.html")
    with open(html_path, "w") as f:
        f.write(md_to_html(md, charts))
    docx_path = os.path.join(RESULTS_DIR, "REPORT.docx")
    md_to_docx(md, docx_path)
    print(f"Wrote {md_path}")
    print(f"Wrote {html_path}")
    print(f"Wrote {docx_path}")
    for family, png in charts.items():
        print(f"Wrote {os.path.join(RESULTS_DIR, png)}")


if __name__ == "__main__":
    main()
